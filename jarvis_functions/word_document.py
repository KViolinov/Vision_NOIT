import os
import time
from docx import Document

from jarvis_functions.essential_functions.voice_input import record_text
from jarvis_functions.essential_functions.enhanced_elevenlabs import generate_audio_from_text

from jarvis_functions.essential_functions.change_config_settings import get_jarvis_voice

jarvis_voice = get_jarvis_voice()

def openWord():
    doc = Document()

    generate_audio_from_text(text="Разбира се, отварям Word. Само секунда", voice=jarvis_voice)

    time.sleep(2)

    generate_audio_from_text(
        text="Готов съм. Преди да започнем, как ще желаете да е заглавието на документа?",
        voice=jarvis_voice)


    print("Listening for title...")
    input_text = record_text()

    doc.add_heading(input_text, 0)

    generate_audio_from_text(
        text="Добре започвам да слушам и записвам. Кажете думата Край за да спра да записвам",
        voice=jarvis_voice)

    words_in_document = ""

    while True:
        print("Listening for input...")
        input_text = record_text()
        print(f"You said: {input_text}")

        if input_text is None or input_text.strip() == "":
            print("No speech detected or input is empty, try again.")
            continue

        if "край" in input_text or "Край" in input_text:
            generate_audio_from_text(text="Спрях да записвам, файла е запазен в папка Downloads",
                                    voice=jarvis_voice)

            doc.add_paragraph(words_in_document)
            break

        words_in_document += input_text + ". "

        time.sleep(1)


    # Finished
    print("Document saved and process ended.")

    file_path = r'D:\example.docx'
    doc.save(file_path)
    os.system(f'start {file_path}')